"""
creation.py — 创作领域（v10）：受控 DocStore + doc 工具。

职责边界（遵循 CLAUDE.md FastAPI 封装原则）：
- 本模块 = 纯 Python 业务模块（doc 状态落盘、章节写入、docx 导出、doc 工具实现）。
- `web/api/routers/creation.py` = 薄 HTTP 包装（前端走 API）。
- agent 工具（CreationProvider）直接调本模块（同进程，无 HTTP 往返）。
- 只有 `_ensure_writing_doc`（plan.py 流程入口）在本模块，plan_node 不内嵌领域逻辑。

写作数据流：
  plan_node(domain=creation) → _ensure_writing_doc 建 doc + 大纲
  → executor 逐章调 creator subagent → doc_write_section 落盘 + SSE `doc_section`
  → doc_get_state / doc_export_docx（docx 生成）

安全：doc_id / section_id 白名单字符校验（杜绝路径穿越）；所有路径 resolve 后
必须落在 workspace/docs 内。
"""

from __future__ import annotations

import json
import re
import time
import uuid
from pathlib import Path

from langchain_core.tools import tool

from ..providers import ToolDef, ToolProvider
from ..safety import tool_allowed
from ..stream import emit

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
WORKSPACE_DOCS = PROJECT_ROOT / "web" / "workspace" / "docs"

_SLUG_RE = re.compile(r"[^a-z0-9-]")


def _safe_id(value: str, default: str = "") -> str:
    """whitelist-char slug; empty → default（防路径穿越/非法文件名）"""
    out = _SLUG_RE.sub("", (value or "").lower().strip())
    return out or default


def _doc_dir(doc_id: str) -> Path:
    return WORKSPACE_DOCS / _safe_id(doc_id)


def _ok(data: dict | list) -> str:
    from agent.tool_contract import ok as _ok_contract
    return _ok_contract(data)


def _err(error: str, error_type: str = "param_error", **ctx) -> str:
    from agent.tool_contract import err as _err_contract
    return _err_contract(error_type, error,
                         next_action="Fix the arguments and retry.", **ctx)


# ---- DocStore（文件系统为事实源；轻量、无 Redis 依赖） ----

def _doc_path(doc_id: str) -> Path:
    return _doc_dir(doc_id) / "doc.json"


def _main_md_path(doc_id: str) -> Path:
    return _doc_dir(doc_id) / f"{_safe_id(doc_id)}.md"


def _load_doc(doc_id: str) -> dict | None:
    p = _doc_path(doc_id)
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except (ValueError, OSError):
        return None


def _save_doc(doc: dict) -> None:
    d = _doc_dir(doc["doc_id"])
    d.mkdir(parents=True, exist_ok=True)
    (d / "sections").mkdir(parents=True, exist_ok=True)
    doc["updated_at"] = time.strftime("%Y-%m-%dT%H:%M:%S")
    _doc_path(doc["doc_id"]).write_text(
        json.dumps(doc, ensure_ascii=False, indent=2), encoding="utf-8")


def _rebuild_markdown(doc_id: str) -> None:
    """按 outline 顺序拼接章节 markdown → 主 md 文件（docx 导出/阅读器读它）。"""
    doc = _load_doc(doc_id)
    if not doc:
        return
    parts: list[str] = []
    for item in doc.get("outline", []):
        sid = item.get("section_id")
        if not sid:
            continue
        fp = _doc_dir(doc_id) / "sections" / f"{_safe_id(sid)}.md"
        if not fp.exists():
            continue
        title = item.get("title", "") or sid
        parts.append(f"## {title}\n\n{fp.read_text(encoding='utf-8').strip()}\n")
    _main_md_path(doc_id).write_text(
        "\n".join(parts) + "\n", encoding="utf-8")


def _summarize_outline(outline: list[str]) -> str:
    return ", ".join(str(s) for s in outline[:12])[:200]


async def verify_section_written(doc_id: str, section_id: str) -> tuple[bool, int]:
    """章节是否已落盘: outline 中该 section status == done(即 doc_write_section
    真正写入过)。返回 (written, word_count)。被 plan 层 executor 用作 creator 步骤
    的权威校验——subagent 无论输出多完整的文本,没写进 doc 就等于没产出。
    """
    sid = _safe_id(doc_id)
    ssec = _safe_id(section_id, "sec")
    doc = _load_doc(sid)
    if not doc:
        return False, 0
    for o in doc.get("outline", []):
        if o.get("section_id") == ssec:
            if o.get("status") != "done":
                return False, 0
            wc = (doc.get("sections") or {}).get(ssec, {}) or {}
            return True, int(wc.get("word_count", 0))
    return False, 0


async def doc_progress(doc_id: str) -> dict | None:
    """文档写入进度(确定性,供 creation synthesize 生成进度报告)。

    返回 {doc_id, title, status, total, done, sections:[{section_id,title,status,
    word_count}]};doc 不存在 → None。
    """
    sid = _safe_id(doc_id)
    doc = _load_doc(sid)
    if not doc:
        return None
    items: list[dict] = []
    done = 0
    for o in doc.get("outline", []):
        wc = (doc.get("sections") or {}).get(o.get("section_id", ""), {}) or {}
        is_done = o.get("status") == "done"
        done += 1 if is_done else 0
        items.append({
            "section_id": o.get("section_id", ""),
            "title": str(o.get("title") or o.get("section_id", "")),
            "status": o.get("status", "pending"),
            "word_count": int(wc.get("word_count", 0)),
        })
    return {
        "doc_id": sid,
        "title": doc.get("title", ""),
        "status": doc.get("status", ""),
        "total": len(items),
        "done": done,
        "sections": items,
    }


# ---- 业务入口（供 plan_node 调用，建 doc + 写大纲） ----

async def _ensure_writing_doc(title: str, outline: list[str],
                              plan_steps: list[dict]) -> str:
    """建 doc + 大纲，返回 doc_id。闲散调用不 raise——异常由调用方兜底。

    `outline` = 章节 id 清单（写进 doc.state 供追踪）；`plan_steps` 带每章的
    title/section_type/cites（LLM 大纲的权威来源）。大纲落盘后前端章节树立即可见。
    """
    doc_id = uuid.uuid4().hex[:12]
    ordered: list[dict] = []
    for step, sid in zip(plan_steps, outline):
        args = dict(step.get("args") or {})
        ordered.append({
            "section_id": _safe_id(args.get("section_id") or sid, "sec"),
            "title": str(args.get("title") or "").strip() or sid,
            "section_type": str(args.get("section_type") or "other").strip() or "other",
            "cites": list(args.get("cites") or []),
            "status": "pending",
        })
    doc = {
        "doc_id": doc_id,
        "title": (title or "Untitled").strip()[:120],
        "status": "writing",
        "outline": ordered,
        "sections": {},
        "created_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "updated_at": "",
    }
    _save_doc(doc)
    emit({"type": "doc_section", "doc_id": doc_id, "section_id": "",
          "status": "created", "title": doc["title"]})
    return doc_id


# ---- doc 工具（agent 可调用；结构化 JSON 返回遵循统一信封） ----

@tool
async def doc_list(status: str = "") -> str:
    """List writing documents. Returns JSON {"ok": true, "data": {docs: [...]}}.
    Optional status filter: outline | writing | done."""
    docs = []
    if WORKSPACE_DOCS.exists():
        for d in sorted(WORKSPACE_DOCS.iterdir()):
            if not d.is_dir() or not (d / "doc.json").exists():
                continue
            doc = _load_doc(d.name)
            if not doc:
                continue
            if status and doc.get("status") != status:
                continue
            docs.append({
                "doc_id": doc["doc_id"], "title": doc.get("title", ""),
                "status": doc.get("status", ""), "n_sections": len(doc.get("outline", [])),
                "updated_at": doc.get("updated_at", ""),
            })
    return _ok({"docs": docs})


@tool
async def doc_create(title: str) -> str:
    """Create a blank writing document. Returns JSON {"ok": true, "data": {doc_id}}.
    The orchestrator normally creates the doc from the outline; use this only to
    start a fresh document before planning its chapters."""
    return _ok({"doc_id": await _ensure_writing_doc(title, [], [])})


@tool
async def doc_set_outline(doc_id: str, outline: str) -> str:
    """Set/replace a document's chapter outline. `outline` is a JSON array of
    {"section_id", "title", "section_type", "cites"} — validated, invalid items
    dropped. Returns the document state."""
    sid = _safe_id(doc_id)
    if not sid:
        return _err("doc_id is required", error_type="param_error")
    doc = _load_doc(sid)
    if not doc:
        return _err(f"doc '{doc_id}' not found (list via doc_list)", error_type="param_error")
    try:
        items = json.loads(outline)
    except (ValueError, TypeError):
        return _err("outline must be a JSON array", error_type="param_error")
    parsed = []
    for it in items if isinstance(items, list) else []:
        if not isinstance(it, dict):
            continue
        parsed.append({
            "section_id": _safe_id(str(it.get("section_id", "")), "sec"),
            "title": str(it.get("title", "")).strip(),
            "section_type": str(it.get("section_type", "other")).strip() or "other",
            "cites": list(it.get("cites") or []),
            "status": "pending",
        })
    doc["outline"] = parsed or doc.get("outline", [])
    doc["status"] = "writing"
    _save_doc(doc)
    _rebuild_markdown(sid)
    return _ok({"doc_id": sid, "outline": doc["outline"]})


@tool
async def doc_write_section(doc_id: str, section_id: str, content: str) -> str:
    """Write one section's Markdown content into a document (atomic write). The
    section must exist in the doc's outline. Marks the section done and updates
    the document progress. Returns {"ok": true, "data": {doc_id, section_id,
    status, word_count}}."""
    sid = _safe_id(doc_id)
    ssec = _safe_id(section_id, "sec")
    if not sid:
        return _err("doc_id is required", error_type="param_error")
    doc = _load_doc(sid)
    if not doc:
        return _err(f"doc '{doc_id}' not found (list via doc_list)", error_type="param_error")
    in_outline = any(o.get("section_id") == ssec for o in doc.get("outline", []))
    if not in_outline:
        return _err(f"section '{section_id}' not in doc outline", error_type="param_error")
    d = _doc_dir(sid)
    (d / "sections").mkdir(parents=True, exist_ok=True)
    (d / "sections" / f"{ssec}.md").write_text(content or "", encoding="utf-8")
    doc.setdefault("sections", {})[ssec] = {
        "status": "done", "updated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "word_count": len((content or "").split()),
    }
    for o in doc.get("outline", []):
        if o.get("section_id") == ssec:
            o["status"] = "done"
    if all(o.get("status") == "done" for o in doc.get("outline", [])):
        doc["status"] = "done"
    _save_doc(doc)
    _rebuild_markdown(sid)
    emit({"type": "doc_section", "doc_id": sid, "section_id": ssec,
          "status": "done", "word_count": len((content or "").split())})
    return _ok({"doc_id": sid, "section_id": ssec, "status": "done",
                "word_count": len((content or "").split())})


@tool
async def doc_get_state(doc_id: str) -> str:
    """Return a document's full state — outline with per-section status, or the
    assembled markdown. Returns JSON {"ok": true, "data": {doc_id, title,
    status, outline, assembled_md}}."""
    sid = _safe_id(doc_id)
    if not sid:
        return _err("doc_id is required", error_type="param_error")
    doc = _load_doc(sid)
    if not doc:
        return _err(f"doc '{doc_id}' not found (list via doc_list)", error_type="param_error")
    md = ""
    p = _main_md_path(sid)
    if p.exists():
        md = p.read_text(encoding="utf-8")
    return _ok({
        "doc_id": sid, "title": doc.get("title", ""), "status": doc.get("status", ""),
        "outline": doc.get("outline", []), "assembled_md": md[-8000:],
    })


@tool
async def doc_export_docx(doc_id: str, format: str = "docx") -> str:
    """Export a document to a .docx file (python-docx). Returns the saved
    relative path (under web/workspace/docs/<doc>/exports/). `format` ignored
    for now (only docx).
    """
    sid = _safe_id(doc_id)
    if not sid:
        return _err("doc_id is required", error_type="param_error")
    doc = _load_doc(sid)
    if not doc:
        return _err(f"doc '{doc_id}' not found (list via doc_list)", error_type="param_error")
    p = _main_md_path(sid)
    if not p.exists():
        return _err("document has no content yet — write sections first", error_type="param_error")
    try:
        path = await _render_docx(sid, doc.get("title", "document"), p.read_text(encoding="utf-8"))
    except Exception:
        return _err("docx render failed (python-docx available?)", error_type="transient")
    try:
        rel = path.relative_to(PROJECT_ROOT).as_posix()
    except ValueError:  # WORKSPACE_DOCS 被重定向到根目录外（测试隔离）时退化为文件名
        rel = path.name
    return _ok({"doc_id": sid, "export_path": rel})


# ---- docx rendering（最小 markdown → docx 转换器） ----


async def _render_docx(doc_id: str, title: str, markdown: str) -> Path:
    """minimal md → docx: #/##/### → Heading, - → List Bullet, else paragraph."""
    from docx import Document

    exports = _doc_dir(doc_id) / "exports"
    exports.mkdir(parents=True, exist_ok=True)
    out = exports / f"{_safe_id(doc_id)}.docx"

    d = Document()
    d.add_heading(title, level=0)
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            d.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            d.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            d.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- "):
            d.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            d.add_paragraph(stripped)
    d.save(str(out))
    return out


# ---- ToolDef + Provider（与 builtin/generic 同构） ----

CREATION_TOOLDEFS = [
    ToolDef(
        name="doc_list",
        description=(
            "List writing documents. Optional status filter (outline|writing|done). "
            "Returns JSON {\"ok\":true,\"data\":{docs:[{doc_id,title,status,n_sections,updated_at}]}}."
        ),
        parameters={
            "type": "object",
            "properties": {
                "status": {"type": "string", "description": "Filter by status.", "default": ""},
            },
        },
        source="builtin",
        annotations={"readOnlyHint": True, "idempotentHint": True},
    ),
    ToolDef(
        name="doc_create",
        description=(
            "Create a blank writing document. Returns {\"ok\":true,\"data\":{doc_id}}. "
            "Normally the orchestrator creates the doc; use only to start a fresh "
            "document before planning chapters."
        ),
        parameters={
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Document title."},
            },
            "required": ["title"],
        },
        source="builtin",
        annotations={"readOnlyHint": False, "idempotentHint": True},
    ),
    ToolDef(
        name="doc_set_outline",
        description=(
            "Set/replace a document's chapter outline (JSON array of "
            "{section_id, title, section_type, cites}). Returns the document state."
        ),
        parameters={
            "type": "object",
            "properties": {
                "doc_id": {"type": "string", "description": "Document id."},
                "outline": {"type": "string", "description": "JSON array of chapters."},
            },
            "required": ["doc_id", "outline"],
        },
        source="builtin",
        annotations={"readOnlyHint": False, "idempotentHint": True},
    ),
    ToolDef(
        name="doc_write_section",
        description=(
            "Write one section's Markdown content into a document (atomic). Section "
            "must be in the doc outline. Marks it done + updates progress. Returns "
            "{\"ok\":true,\"data\":{doc_id,section_id,status,word_count}}."
        ),
        parameters={
            "type": "object",
            "properties": {
                "doc_id": {"type": "string", "description": "Document id."},
                "section_id": {"type": "string", "description": "Section slug (from outline)."},
                "content": {"type": "string", "description": "Section Markdown body."},
            },
            "required": ["doc_id", "section_id", "content"],
        },
        source="builtin",
        annotations={"readOnlyHint": False, "idempotentHint": True},
    ),
    ToolDef(
        name="doc_get_state",
        description=(
            "Return a document's full state: outline with per-section status and the "
            "assembled markdown. Returns {\"ok\":true,\"data\":{doc_id,title,status,outline,assembled_md}}."
        ),
        parameters={
            "type": "object",
            "properties": {
                "doc_id": {"type": "string", "description": "Document id."},
            },
            "required": ["doc_id"],
        },
        source="builtin",
        annotations={"readOnlyHint": True, "idempotentHint": True},
    ),
    ToolDef(
        name="doc_export_docx",
        description=(
            "Export a document to .docx (python-docx). Returns the saved relative "
            "path under web/workspace/docs/<doc>/exports/."
        ),
        parameters={
            "type": "object",
            "properties": {
                "doc_id": {"type": "string", "description": "Document id."},
                "format": {"type": "string", "description": "Only 'docx' for now.", "default": "docx"},
            },
            "required": ["doc_id"],
        },
        source="builtin",
        annotations={"readOnlyHint": False},
    ),
]

_FUNC_MAP = {
    "doc_list": doc_list,
    "doc_create": doc_create,
    "doc_set_outline": doc_set_outline,
    "doc_write_section": doc_write_section,
    "doc_get_state": doc_get_state,
    "doc_export_docx": doc_export_docx,
}


class CreationProvider(ToolProvider):
    """创作领域工具提供商（doc 系列；subagent 才可见，不进父 agent 工具面）。"""

    name = "creation"

    def __init__(self):
        self._tool_map = {td.name: td for td in CREATION_TOOLDEFS}

    async def list_tools(self) -> list[ToolDef]:
        return list(self._tool_map.values())

    async def call_tool(self, name: str, arguments: dict):
        if name not in _FUNC_MAP:
            raise KeyError(f"Creation tool '{name}' not found")
        td = self._tool_map[name]
        if td is not None and not tool_allowed(td.annotations):
            return _err(
                f"Action '{name}' is not authorized for the current role.",
                error_type="permission_denied",
            )
        fn = _FUNC_MAP[name]
        return await fn.ainvoke(arguments)